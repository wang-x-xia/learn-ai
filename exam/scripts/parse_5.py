# /// script
# requires-python = ">=3.12"
# dependencies = [
#     "python-docx",
# ]
# ///

from docx import Document
import re
import json

def convert_tf_answer(answer):
    """将判断题答案转换为 bool"""
    if answer == '√':
        return True
    elif answer == '×':
        return False
    return answer

def convert_multi_answer(answer):
    """将多选题答案转换为字母数组"""
    if isinstance(answer, str) and len(answer) > 1:
        return list(answer)
    return answer

def parse_exam_answers(file_path):
    """解析模拟试卷的答案"""
    doc = Document(file_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    answers = {
        '判断题': {},
        '单选题': {},
        '多选题': {}
    }

    current_type = None

    for i, para in enumerate(paragraphs):
        # 检测答案部分开始
        if '判断题参考答案' in para or '判断题答案' in para:
            current_type = '判断题'
            print(f"找到判断题答案 at line {i}")
            continue
        elif '单选题参考答案' in para or '单选题答案' in para:
            current_type = '单选题'
            print(f"找到单选题答案 at line {i}")
            continue
        elif '多选题参考答案' in para or '多选题答案' in para:
            current_type = '多选题'
            print(f"找到多选题答案 at line {i}")
            continue

        # 解析答案行
        if current_type:
            # 判断题特殊格式：1-10）(√) (×) (×) (√) (×) (×) (×) (√) (×) (√)
            if current_type == '判断题':
                tf_match = re.match(r'(\d+)-(\d+)[）\)]\s*(.+)', para)
                if tf_match:
                    start = int(tf_match.group(1))
                    end = int(tf_match.group(2))
                    answer_part = tf_match.group(3)
                    # 提取所有 (√) 或 (×)
                    tf_answers = re.findall(r'[（(]\s*([×√])\s*[）)]', answer_part)
                    for j, ans in enumerate(tf_answers):
                        if j < (end - start + 1):
                            answers[current_type][start + j] = convert_tf_answer(ans)
                    continue

            # 选择题格式：1-10) ABDAAAADDD 或 1）A 或 1.A（支持A-E）
            range_match = re.search(r'(\d+)-(\d+)[）\)]\s*([A-E]+)', para)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2))
                answer_str = range_match.group(3)
                for j, char in enumerate(answer_str):
                    if j < (end - start + 1):
                        if current_type == '多选题':
                            answers[current_type][start + j] = convert_multi_answer(char)
                        else:
                            answers[current_type][start + j] = char
                continue

            # 单个题目答案：1）A 或 1.A 或 1) ABC（支持多选题的多个字母，支持A-E）
            single_match = re.match(r'(\d+)[）.)]\s*([A-E]+)', para)
            if single_match:
                num = int(single_match.group(1))
                answer = single_match.group(2)
                if current_type == '多选题':
                    answers[current_type][num] = convert_multi_answer(answer)
                else:
                    answers[current_type][num] = answer
                continue

    return answers

def parse_questions(file_path):
    doc = Document(file_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    
    questions = {
        '判断题': [],
        '单选题': [],
        '多选题': []
    }
    
    current_type = None
    current_question = None
    
    for i, para in enumerate(paragraphs):
        # 检测题型
        if '判断题' in para and '一、' in para:
            # 先保存当前题目（如果有）
            if current_question:
                questions[current_type].append(current_question)
                current_question = None
            current_type = '判断题'
            print(f"切换到题型: {current_type} at line {i}")
            continue
        elif '单选题' in para:
            # 先保存当前题目（如果有）
            if current_question:
                questions[current_type].append(current_question)
                current_question = None
            current_type = '单选题'
            print(f"切换到题型: {current_type} at line {i}")
            continue
        elif '多选题' in para:
            # 先保存当前题目（如果有）
            if current_question:
                questions[current_type].append(current_question)
                current_question = None
            current_type = '多选题'
            print(f"切换到题型: {current_type} at line {i}")
            continue
        
        # 检测题目编号
        # 判断题: （    ）1. 或 (    ) 1.
        if re.match(r'^[（(]\s*[）)]\s*\d+\.', para):
            if current_question:
                questions[current_type].append(current_question)
            # 使用当前题型长度+1作为新ID
            new_id = len(questions[current_type]) + 1
            # 判断题不需要 options 字段
            current_question = {
                'id': new_id,
                'question': para,
                'answer': None
            }
        # 单选题/多选题: 1. 或 (1) 或 1、
        elif re.match(r'^\d+\.', para) or re.match(r'^[（(]\d+[）)]\s*', para):
            # 检查是否是选项行（支持A-E）
            if not re.match(r'^[（(]\s*[A-E]\s*[）)]', para) and not re.match(r'^[A-E][、.]', para):
                if current_question:
                    questions[current_type].append(current_question)
                # 使用当前题型长度+1作为新ID
                new_id = len(questions[current_type]) + 1
                current_question = {
                    'id': new_id,
                    'question': para,
                    'options': [],
                    'answer': None
                }
        # 检测选项 (A) 或 A、（支持A-E）
        elif re.match(r'^[（(]\s*[A-E]\s*[）)]', para) or re.match(r'^[A-E][、.]', para):
            if current_question:
                current_question['options'].append(para)
    
    # 添加最后一个问题
    if current_question:
        questions[current_type].append(current_question)
    
    return questions

def show_sample(questions, qtype, count=2):
    print(f"\n{qtype} 示例 (前{count}题):")
    for i, q in enumerate(questions[qtype][:count]):
        print(f"\n题目 {q['id']}:")
        print(f"  问题: {q['question']}")
        if 'options' in q and q['options']:
            print(f"  选项:")
            for opt in q['options']:
                print(f"    {opt}")
        else:
            print(f"  选项: 无")

# 处理模拟试卷
print(f"{'='*60}")
print("处理理论知识模拟试卷")
print(f"{'='*60}")
questions = parse_questions("../source/第5部分_人工智能训练师_3级_理论知识模拟试卷.docx")

print(f"\n题目统计:")
for qtype, qlist in questions.items():
    print(f"  {qtype}: {len(qlist)} 题")

# 解析答案
print(f"\n{'='*60}")
print("解析模拟试卷答案")
print(f"{'='*60}")
answers = parse_exam_answers("../source/第5部分_人工智能训练师_3级_理论知识模拟试卷.docx")

# 手动添加单选题1-10的答案
for i, char in enumerate('ABDAAAADDD'):
    answers['单选题'][i + 1] = char

print(f"\n答案统计:")
for atype, adict in answers.items():
    print(f"  {atype}: {len(adict)} 个答案")

# 合并答案到题目
for qtype in ['判断题', '单选题', '多选题']:
    for q in questions[qtype]:
        q_id = q['id']
        # 尝试用字符串ID和数字ID匹配
        if str(q_id) in answers[qtype]:
            q['answer'] = answers[qtype][str(q_id)]
        elif q_id in answers[qtype]:
            q['answer'] = answers[qtype][q_id]
        else:
            # 如果没有匹配的答案，保持None
            pass

# 后处理：确保多选题答案是数组格式
for q in questions['多选题']:
    if q['answer'] and isinstance(q['answer'], str):
        q['answer'] = convert_multi_answer(q['answer'])

# 统计答案匹配情况
print(f"\n答案匹配统计:")
for qtype in ['判断题', '单选题', '多选题']:
    total = len(questions[qtype])
    with_answer = sum(1 for q in questions[qtype] if q['answer'])
    print(f"  {qtype}: {total} 题，有答案: {with_answer} 题")

show_sample(questions, '判断题')
show_sample(questions, '单选题')
show_sample(questions, '多选题')

# 保存
with open('../data/questions_5.json', 'w', encoding='utf-8') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)

print(f"\n已保存到 ../data/questions_5.json")
